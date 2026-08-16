"""Renders report/report.md -> report/report.docx (Word), so the report can be
reviewed/edited/exported to PDF by hand before submission. Custom lightweight
Markdown -> docx renderer (tailored to this project's own report.md structure:
headers, paragraphs with **bold**/*italic*, pipe tables, bullet/numbered
lists, blockquotes, and ![alt](path) images) rather than a generic converter,
since pandoc isn't available on this machine.

Run: .venv/Scripts/python.exe -m code.utils.build_report_docx
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

_ROOT = Path(__file__).resolve().parents[2]
_MD_PATH = _ROOT / "report" / "report.md"
_DOCX_PATH = _ROOT / "report" / "report.docx"

_INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*|`.+?`)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Splits on **bold**, *italic*, `code` and adds runs preserving order."""
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            paragraph.add_run(part)


def _parse_table(lines: list[str], start_idx: int):
    """Parses a pipe-table starting at lines[start_idx]; returns (rows, next_idx)."""
    rows = []
    i = start_idx
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|[\s:|-]+\|$", line):  # separator row
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


def build_docx(md_path: Path = _MD_PATH, docx_path: Path = _DOCX_PATH) -> Path:
    text = md_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("<!--"):
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped.lstrip("#").strip()
            level = min(max(level, 1), 4)
            h = doc.add_heading(level=level)
            _add_inline_runs(h, heading_text)
            i += 1
            continue

        if stripped.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if m:
                alt, img_path = m.group(1), m.group(2)
                # Image paths in report.md are relative to report.md's OWN
                # directory (e.g. "../figures/fig01....png"), not the project
                # root -- resolve against md_path.parent, matching how any
                # normal Markdown renderer (and GitHub) would interpret them.
                full_path = (md_path.parent / img_path).resolve()
                if full_path.exists():
                    doc.add_picture(str(full_path), width=Cm(15))
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.add_run(alt)
                    run.italic = True
                    run.font.size = Pt(9)
                else:
                    doc.add_paragraph(f"[MISSING FIGURE: {img_path}]")
            i += 1
            continue

        if stripped.startswith("|"):
            rows, next_i = _parse_table(lines, i)
            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(n_cols):
                        cell_text = row[c_idx] if c_idx < len(row) else ""
                        cell = table.cell(r_idx, c_idx)
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _add_inline_runs(p, cell_text)
                        if r_idx == 0:
                            for run in p.runs:
                                run.bold = True
            i = next_i
            continue

        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            _add_inline_runs(p, stripped.lstrip(">").strip())
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Regular paragraph (accumulate until blank line for readable wrapping)
        para_lines = [stripped]
        j = i + 1
        while j < n and lines[j].strip() and not lines[j].strip().startswith(("#", "|", ">", "!["))\
                and not re.match(r"^[-*]\s+", lines[j].strip()) and not re.match(r"^\d+\.\s+", lines[j].strip()):
            para_lines.append(lines[j].strip())
            j += 1
        p = doc.add_paragraph()
        _add_inline_runs(p, " ".join(para_lines))
        i = j

    doc.save(str(docx_path))
    return docx_path


if __name__ == "__main__":
    out = build_docx()
    print(f"Wrote {out} ({out.stat().st_size / 1024:.1f} KB)")
